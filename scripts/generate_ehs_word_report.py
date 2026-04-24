import json
import os
from datetime import date
from docx import Document
from docx.shared import Inches


def role_workflow(role):
    flows = {
        "Admin": [
            "Log in to EHS tenant context.",
            "Review dashboard summary and operational status.",
            "Access administrative modules for users/settings/governance.",
            "Validate cross-module readiness for daily operations."
        ],
        "Nurse": [
            "Log in to EHS and access patient-facing workflows.",
            "Review queue and patient records from allowed modules.",
            "Update clinical workflow data permitted to nursing role.",
            "Coordinate with doctor/lab and continue care pathway."
        ],
        "Lab": [
            "Log in to EHS and open lab-relevant modules.",
            "Review diagnostic workload and patient test context.",
            "Record and update test results in assigned workflow areas.",
            "Return control to clinician-facing workflow."
        ],
        "Insurance": [
            "Log in and access insurance registry/claims modules.",
            "Review provider and claim records.",
            "Create or update claim lifecycle status.",
            "Share status with billing/accounts stakeholders."
        ],
        "Billing": [
            "Log in and open billing-related modules.",
            "Review invoice queue and payment states.",
            "Issue invoices and update settlement status.",
            "Coordinate with accounts for reconciliation."
        ],
        "Accounts": [
            "Log in and open accounts/financial modules.",
            "Review financial snapshot and expense flow.",
            "Track period financial performance and outflow.",
            "Report summary status to management."
        ],
        "Auditor": [
            "Log in with read-focused auditing context.",
            "Review report and compliance-relevant modules.",
            "Validate access boundaries and transactional traceability.",
            "Raise findings for governance follow-up."
        ],
        "Management": [
            "Log in and review executive dashboard metrics.",
            "Access reports and strategic operational views.",
            "Validate performance across clinical and financial tracks.",
            "Drive high-level decisions and escalations."
        ],
    }
    return flows.get(role, [
        "Log in to EHS tenant context.",
        "Validate role-scoped module access.",
        "Execute assigned operational workflow.",
        "Confirm expected navigation and system behavior."
    ])


def main():
    repo_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    today = date.today().isoformat()
    base_dir = os.path.join(repo_root, "Docs", "07-Deployment", f"ehs-role-snapshots-{today}")
    json_path = os.path.join(base_dir, "ehs_role_capture_results.json")

    if not os.path.exists(json_path):
        raise FileNotFoundError(f"Capture result not found: {json_path}")

    with open(json_path, "r", encoding="utf-8") as f:
        payload = json.load(f)

    doc = Document()
    doc.add_heading("EHS Role Access Review and Workflow Evidence", 0)
    doc.add_paragraph(f"Date: {payload.get('date', today)}")
    doc.add_paragraph("Tenant Scope: Enterprise Hospital Systems (EHS) only")
    doc.add_paragraph("Purpose: User-facing pre-delivery document with architecture context, role workflows, and screenshot evidence.")

    doc.add_heading("1. Technology Stack", level=1)
    stack_points = [
        "Frontend: React 18 with Vite build/runtime.",
        "Backend: Node.js + Express REST API.",
        "Database: PostgreSQL (shared schema with tenant isolation).",
        "Authentication: JWT + bcrypt password hashing.",
        "Authorization: Role-permission model + module/feature gating.",
        "Testing: Playwright E2E/smoke + integration scripts."
    ]
    for item in stack_points:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("2. Architecture Summary", level=1)
    arch_points = [
        "Application pattern: Multi-tenant SPA + stateless API backend.",
        "Core frontend orchestrator: client/src/App.jsx.",
        "API routing: server/index.js.",
        "Security middleware: authenticate, requireTenant, requirePermission.",
        "Data access layer: server/db/repository.js using parameterized SQL.",
        "Tenant boundary: tenant_id scoping in middleware and repository access."
    ]
    for item in arch_points:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3. EHS Snapshot Execution Summary", level=1)
    doc.add_paragraph(f"Base URL: {payload.get('baseUrl')}")
    doc.add_paragraph(f"Roles attempted: {payload.get('total')}")
    doc.add_paragraph(f"Successful captures: {payload.get('success')}")
    doc.add_paragraph(f"Failed captures: {payload.get('failed')}")

    doc.add_heading("4. Role-wise Workflow and Evidence", level=1)

    for entry in payload.get("results", []):
        role = entry.get("role", "Unknown")
        doc.add_heading(f"{role} - {entry.get('tenant', 'EHS')}", level=2)
        doc.add_paragraph(f"User: {entry.get('email', '-')}")
        doc.add_paragraph(f"Login status: {entry.get('status', '-')}")

        nav_items = entry.get("navItems", [])
        if nav_items:
            doc.add_paragraph("Visible navigation:", style="List Bullet")
            for nav in nav_items:
                doc.add_paragraph(nav, style="List Bullet 2")
        elif entry.get("status") == "failed":
            doc.add_paragraph(f"Error: {entry.get('error', 'Unknown')}")

        doc.add_paragraph("Workflow steps:")
        for step in role_workflow(role):
            doc.add_paragraph(step, style="List Number")

        image_name = entry.get("screenshot")
        if image_name:
            image_path = os.path.join(base_dir, image_name)
            if os.path.exists(image_path):
                doc.add_paragraph("Screenshot evidence:")
                doc.add_picture(image_path, width=Inches(6.5))
            else:
                doc.add_paragraph(f"Screenshot missing: {image_name}")

    out_docx = os.path.join(repo_root, "Docs", "07-Deployment", f"EHS_Role_Workflow_Report_{today}.docx")
    doc.save(out_docx)
    print(f"Saved: {out_docx}")


if __name__ == "__main__":
    main()
